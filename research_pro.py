import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from collections import Counter
import re
import matplotlib as mpl
from wordcloud import WordCloud
from io import BytesIO
from docx import Document
import networkx as nx 
from langdetect import detect # ระบบตรวจจับภาษาอัตโนมัติ

# --- 1. ตั้งค่าการแสดงผลภาษาไทย ---
font_path = "Kanit-Regular.ttf" 

def setup_font():
    try:
        mpl.font_manager.fontManager.addfont(font_path)
        prop = mpl.font_manager.FontProperties(fname=font_path)
        mpl.rc('font', family=prop.get_name(), size=12)
        mpl.rcParams['axes.unicode_minus'] = False 
        return prop
    except:
        return None

# --- 2. ฟังก์ชันช่วยสร้างไฟล์ Export ---
def to_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Summary')
    return output.getvalue()

def create_word_report(filename, sentiment, summary, keywords_df, original_text):
    doc = Document()
    doc.add_heading(f'Research Report: {filename}', 0)
    doc.add_heading('1. Sentiment Analysis', level=1)
    doc.add_paragraph(sentiment)
    doc.add_heading('2. Executive Summary', level=1)
    for s in summary:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_heading('3. Key Findings (Top Keywords)', level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Keyword'
    hdr_cells[1].text = 'Frequency'
    for index, row in keywords_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row.iloc[0])
        row_cells[1].text = str(row.iloc[1])
    doc.add_heading('4. Source Text Snippet', level=1)
    doc.add_paragraph(original_text[:1000] + "..." if len(original_text) > 1000 else original_text)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 3. ฟังก์ชันวิเคราะห์อารมณ์ (Dual Language) ---
def analyze_sentiment(text, lang):
    if lang == 'th':
        pos = ['ดี', 'สำเร็จ', 'ภูมิใจ', 'ความสุข', 'พัฒนา', 'ประโยชน์', 'ยั่งยืน', 'พอเพียง', 'ประหยัด']
        neg = ['ไม่ดี', 'ปัญหา', 'แย่', 'ยากลำบาก', 'ขาดแคลน', 'อุปสรรค', 'หนี้สิน', 'เดือดร้อน']
    else:
        pos = ['good', 'great', 'excellent', 'success', 'happy', 'positive', 'improve', 'benefit', 'sustainable']
        neg = ['bad', 'problem', 'difficult', 'lack', 'obstacle', 'debt', 'negative', 'poor', 'issue']
    
    pos_score = sum(1 for w in pos if w in text.lower())
    neg_score = sum(1 for w in neg if w in text.lower())
    
    if pos_score > neg_score: return "บวก (Positive) 😊"
    elif neg_score > pos_score: return "ลบ (Negative) 😟"
    else: return "ปกติ / เป็นกลาง 😐"

# --- 4. ฟังก์ชันสร้าง Network Analysis ---
def plot_network(words, font_prop):
    G = nx.Graph()
    pairs = [tuple(sorted((words[i], words[i+1]))) for i in range(len(words)-1) if words[i] != words[i+1]]
    pair_counts = Counter(pairs).most_common(20)
    for pair, weight in pair_counts:
        G.add_edge(pair[0], pair[1], weight=weight)
    if len(G.nodes) == 0: return None
    
    fig, ax = plt.subplots(figsize=(10, 7))
    pos = nx.spring_layout(G, k=0.5, seed=42)
    weights = [G[u][v]['weight'] for u, v in G.edges()]
    nx.draw_networkx_edges(G, pos, width=weights, edge_color='skyblue', alpha=0.5)
    nx.draw_networkx_nodes(G, pos, node_size=2000, node_color='orange', alpha=0.8)
    
    for node, (x, y) in pos.items():
        ax.text(x, y, node, fontproperties=font_prop, fontsize=12, ha='center', va='center', 
                bbox=dict(facecolor='white', alpha=0.6, edgecolor='none'))
    plt.axis('off')
    return fig

# --- 5. เริ่มต้นโปรแกรม ---
try:
    from pythainlp.tokenize import word_tokenize
    from pythainlp.corpus import thai_stopwords
    from pythainlp.summarize import summarize
    THAI_READY = True
except:
    THAI_READY = False

st.set_page_config(layout="wide", page_title="Professional Research Analysis Tool")
st.title("🌍 ระบบวิเคราะห์งานวิจัยขั้นสูง (Multi-Language Pro)")

if not THAI_READY:
    st.error("❌ Library ไม่พร้อมใช้งาน กรุณาตรวจสอบ requirements.txt")
    st.stop()

font_p = setup_font()
uploaded_files = st.file_uploader("อัปโหลดไฟล์บทสัมภาษณ์หรือเอกสาร (.txt)", type=['txt'], accept_multiple_files=True)

if uploaded_files:
    comparison_list = []
    
    for file in uploaded_files:
        raw_text = file.read().decode("utf-8")
        
        # ตรวจจับภาษา
        try:
            lang = detect(raw_text)
        except:
            lang = 'th'
        
        # ตั้งค่าตามภาษา
        if lang == 'th':
            tokens = word_tokenize(raw_text, keep_whitespace=False)
            stop_words = list(thai_stopwords()) + ['เนาะ', 'นะ', 'ครับ', 'ค่ะ', 'คือ', 'แบบ', 'ว่า']
            min_len = 5
            display_lang = "ภาษาไทย 🇹🇭"
        else:
            tokens = re.findall(r'\w+', raw_text.lower())
            stop_words = ['the', 'and', 'is', 'in', 'to', 'of', 'for', 'with', 'it', 'on', 'as', 'that', 'this', 'was', 'were', 'have', 'has']
            min_len = 3
            display_lang = "English 🇺🇸"

        filtered_words = [t.strip() for t in tokens if t.strip() and t.lower() not in stop_words and len(t.strip()) >= min_len and not re.match(r'^[0-9\W]+$', t)]
        word_counts = Counter(filtered_words)
        filtered_final = [w for w in filtered_words if word_counts[w] >= 3]
        
        s_label = analyze_sentiment(raw_text, lang)
        try:
            brief = summarize(raw_text, n=2)
        except:
            brief = ["ไม่สามารถสรุปเนื้อหาได้"]

        comparison_list.append({
            "ชื่อไฟล์": file.name,
            "ภาษา": display_lang,
            "โทนอารมณ์": s_label,
            "คำสำคัญเด่น": Counter(filtered_final).most_common(1)[0][0] if filtered_final else "N/A"
        })

        with st.expander(f"📑 ผลการวิเคราะห์ละเอียด: {file.name} ({display_lang})", expanded=True):
            tab1, tab2, tab3 = st.tabs(["📊 สถิติ & Word Cloud", "🕸️ โครงข่ายความสัมพันธ์", "📄 ต้นฉบับ"])
            
            with tab1:
                col_a, col_b = st.columns(2)
                with col_a:
                    st.subheader("💡 การวิเคราะห์เนื้อหา")
                    st.write(f"**ความรู้สึก:** {s_label}")
                    st.write("**สรุปประเด็น:**")
                    for b in brief: st.write(f"📌 {b}")
                    
                    if filtered_final:
                        wc = WordCloud(width=800, height=400, background_color="white", regexp=r"[\w\u0e00-\u0e7f]+", font_path=font_path).generate(" ".join(filtered_final))
                        fig_wc, ax_wc = plt.subplots()
                        ax_wc.imshow(wc, interpolation='bilinear')
                        ax_wc.axis("off")
                        st.pyplot(fig_wc)
                        
                        buf_wc = BytesIO()
                        fig_wc.savefig(buf_wc, format="png")
                        st.download_button("💾 โหลด Word Cloud (PNG)", buf_wc.getvalue(), f"cloud_{file.name}.png", "image/png")
                
                with col_b:
                    st.subheader("📈 สถิติคำสำคัญ")
                    df_counts = pd.DataFrame(Counter(filtered_final).most_common(12), columns=['คำ', 'จำนวนครั้ง'])
                    st.table(df_counts)
                    st.download_button("🟢 โหลดสถิตินี้ (Excel)", to_excel(df_counts), f"stats_{file.name}.xlsx")
                    
                    word_report = create_word_report(file.name, s_label, brief, df_counts, raw_text)
                    st.download_button("📄 โหลดรายงานสรุป (Word)", word_report, f"report_{file.name}.docx")

            with tab2:
                st.subheader("🕸️ ความสัมพันธ์ระหว่างคำ (Network Analysis)")
                if len(filtered_words) > 5:
                    fig_net = plot_network(filtered_words, font_p)
                    if fig_net:
                        st.pyplot(fig_net)
                        buf_net = BytesIO()
                        fig_net.savefig(buf_net, format="png")
                        st.download_button("💾 โหลดภาพโครงข่าย (PNG)", buf_net.getvalue(), f"network_{file.name}.png", "image/png")
                else:
                    st.warning("ข้อมูลไม่เพียงพอสำหรับสร้างโครงข่าย")

            with tab3:
                st.subheader("📄 ตัวอย่างข้อความและไฟล์ต้นฉบับ")
                sentences = [s.strip() for s in raw_text.split('\n') if s.strip()]
                st.info("\n\n".join(sentences[:5]))
                st.divider()
                st.text_area("Full Content", value=raw_text, height=300)

    # --- ส่วนสรุปภาพรวมทั้งหมด ---
    st.divider()
    st.subheader("📋 ตารางเปรียบเทียบทุกเคสในโปรเจกต์")
    df_compare = pd.DataFrame(comparison_list)
    st.table(df_compare)
    st.download_button("🟢 ดาวน์โหลดตารางรวม (Excel)", to_excel(df_compare), "global_summary.xlsx")

else:
    st.info("👋 ยินดีต้อนรับ! กรุณาอัปโหลดไฟล์ .txt เพื่อเริ่มการวิเคราะห์เชิงลึกทั้งไทยและอังกฤษ")
